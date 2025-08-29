from routes import main_blueprint
from .common_imports import *
from routes.routes_public import safe_title

def sanitize_filename(filename):
    """Replaces spaces and invalid characters with underscores."""
    # Replace spaces with underscores first
    filename = filename.replace(" ", "_")
    # Replace any character that is not alphanumeric, underscore, or dot with an underscore
    filename = re.sub(r'[^\w.]+', '_', filename)
    # Optional: Collapse multiple underscores into one
    filename = re.sub(r'_+', '_', filename)
    return filename

def save_file(file, name=None):
    """Securely saves an uploaded file with a short unique name."""
    if name:
        filename = secure_filename(name)
    else:
        filename = secure_filename(file.filename)

    # Sanitize the filename
    sanitized_filename = sanitize_filename(filename)

    # Use shorter unique id: 6-digit microsecond timestamp
    timestamp = datetime.now().strftime('%H%M%S%f')[:9] # Use microseconds for better uniqueness
    unique_filename = f"{timestamp}_{filename}"
    file_path = os.path.join("static/uploads", unique_filename)

    file.save(file_path)
    return unique_filename, file_path



def download_image(image_url):
    """Downloads an image from a URL, saves it locally with a short unique name, and returns filename & path."""
    try:
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0 Safari/537.36"
            )
        }
        response = requests.get(image_url, headers=headers, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        raise Exception(f"Failed to download image: {e}")

    # Derive safe base name
    parsed = urlparse(image_url)
    name = os.path.basename(parsed.path) or "DownImg"
    base_filename = secure_filename(name)

    # Sanitize the base filename
    sanitized_base_filename = sanitize_filename(base_filename)

    # Ensure extension exists
    if not os.path.splitext(sanitized_base_filename)[1]:
        ext = guess_extension(response.headers.get("Content-Type", "").split(";")[0])
        sanitized_base_filename += ext or ".jpg"

    # Short timestamp
    timestamp = datetime.now().strftime('%H%M%S%f')[:9] # Use microseconds
    unique_filename = f"{timestamp}_{base_filename}"

    file_path = os.path.join("static/uploads", unique_filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    with open(file_path, "wb") as f:
        f.write(response.content)

    return unique_filename, file_path



def delete_existing_file(file_path):
    """Deletes an existing file from storage if it exists."""
    if file_path and os.path.exists(file_path):
        os.remove(file_path)


def parse_iso8601_duration(duration):
    """Convert ISO8601 duration (e.g. PT1H2M3S) to seconds."""
    pattern = re.compile(
        r"^PT(?:(?P<hours>\d+)H)?(?:(?P<minutes>\d+)M)?(?:(?P<seconds>\d+)S)?$"
    )
    m = pattern.match(duration)
    if not m:
        return None
    p = m.groupdict()
    return (
        int(p.get("hours") or 0) * 3600
        + int(p.get("minutes") or 0) * 60
        + int(p.get("seconds") or 0)
    )

def process_single_image(file_path):
    """Standardizes an image to WebP (fixed width 1280px, preserve aspect ratio), creates 3 types of thumbnails, deletes original."""
    try:
        img = PILImage.open(file_path).convert("RGB")
        base_dir = os.path.dirname(file_path)
        base_name, _ = os.path.splitext(os.path.basename(file_path))
        webp_path = os.path.join(base_dir, f"{base_name}.webp")

        # 📏 Standard Width
        STANDARD_WIDTH = 1280

        # Resize keeping aspect ratio
        w_percent = STANDARD_WIDTH / float(img.size[0])
        new_height = int(float(img.size[1]) * float(w_percent))
        resized_img = img.resize((STANDARD_WIDTH, new_height), PILImage.Resampling.LANCZOS)

        # Save standardized WebP
        resized_img.save(webp_path, "WEBP", quality=85, method=6)

        # 📸 Create Thumbnails
        THUMB_SIZES = {
            "portrait": (400, 600),   # 2:3
            "square": (400, 400),     # 1:1
            "landscape": (600, 400),  # 3:2
        }

        for style, size in THUMB_SIZES.items():
            thumb = img.copy()
            thumb = ImageOps.fit(img, size, PILImage.Resampling.LANCZOS, centering=(0.5, 0.5))
            thumb_path = os.path.join(base_dir, f"{base_name}_thumb_{style}.webp")
            thumb.save(thumb_path, "WEBP", quality=75, method=6)

        # ✅ Delete the original file
        if not file_path.endswith(".webp"):
            os.remove(file_path)

        return "ok"

    except Exception as e:
        return f"Processing error: {str(e)}"


@main_blueprint.route("/admin/tools/docx-upload")
@login_required
def tools_docx_upload():
    """Admin tool to upload DOCX files and convert them to News articles."""
    if not current_user.verified:
        flash("Akun Anda harus diverifikasi untuk mengakses tools ini.", "error")
        return redirect(url_for("main.admin_dashboard"))
    
    # Get categories for the form
    categories = Category.query.order_by(Category.name).all()
    
    # Get today's date for default value
    from datetime import date
    today = date.today().isoformat()
    
    return render_template(
                    "admin/tools/docx_uploader/index.html",
        categories=categories,
        today=today,
        title="Upload DOCX ke News"
    )


@main_blueprint.route("/admin/tools/docx-template", methods=["GET"])
@login_required
def download_docx_template():
    """Serve a downloadable DOCX template for writers (Indonesian)."""
    if not current_user.verified:
        flash("Akun Anda harus diverifikasi untuk mengakses tools ini.", "error")
        return redirect(url_for("main.admin_dashboard"))

    from docx import Document
    from docx.shared import Pt, Inches
    from io import BytesIO

    doc = Document()

    # Title style
    title = doc.add_heading("Judul Artikel (H1)", level=1)

    # Metadata guidance
    doc.add_paragraph("Ringkasan singkat (opsional, 1-2 kalimat) untuk meta deskripsi.")
    doc.add_paragraph("Tag/Kata Kunci (opsional): pisahkan dengan koma, maksimal 5.")

    # Content sections
    doc.add_heading("Subjudul Utama (H2)", level=2)
    doc.add_paragraph(
        "Tulis paragraf pembuka di sini. Gunakan paragraf pendek, jelas, dan mudah dibaca.")

    doc.add_heading("Subjudul Tambahan (H2)", level=2)
    doc.add_paragraph("Gunakan daftar jika diperlukan:")
    ul = doc.add_paragraph()
    ul.style = 'List Bullet'
    ul.add_run(" Poin penting 1")
    doc.add_paragraph(" Poin penting 2", style='List Bullet')

    doc.add_paragraph("Tambahkan kutipan jika diperlukan:")
    quote = doc.add_paragraph()
    quote.style = 'Intense Quote'
    quote.add_run("\"Kutipan penting terkait topik.\"")

    doc.add_heading("Gambar (Opsional)", level=2)
    doc.add_paragraph(
        "Anda dapat menyisipkan gambar di sini. Pastikan mencantumkan sumber jika diambil dari pihak ketiga.")

    doc.add_heading("Catatan untuk Penulis", level=3)
    doc.add_paragraph(
        "- Gunakan bahasa Indonesia yang baik dan benar.\n"
        "- Hindari kalimat terlalu panjang.\n"
        "- Gunakan subjudul untuk memecah topik.\n"
        "- Sertakan sumber/referensi jika mengutip data.\n"
        "- Jika konten sensitif (dewasa), mohon beri catatan di awal dokumen.")

    # Footer guidance
    doc.add_paragraph()
    doc.add_paragraph(
        "Template ini akan diproses otomatis menjadi artikel News. Judul H1 akan menjadi judul artikel. "
        "Gunakan H2/H3 untuk struktur yang jelas.")

    # Prepare response
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="Template_Artikel_News.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@main_blueprint.route("/admin/tools/docx-contoh", methods=["GET"])
@login_required
def download_docx_contoh():
    """Serve the example DOCX file for writers to see how the format should look."""
    if not current_user.verified:
        flash("Akun Anda harus diverifikasi untuk mengakses tools ini.", "error")
        return redirect(url_for("main.admin_dashboard"))

    # Path to the example DOCX file
    contoh_path = os.path.join(current_app.root_path, "templates", "admin", "tools", "docx_uploader", "contoh.docx")
    
    if not os.path.exists(contoh_path):
        flash("File contoh tidak ditemukan.", "error")
        return redirect(url_for("main.tools_docx_upload"))

    return send_file(
        contoh_path,
        as_attachment=True,
        download_name="Contoh_Artikel_News.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# Helper function (keep as before)
def add_url_to_sitemap(urlset, loc, lastmod=None, changefreq=None, priority=None):
    """Adds a <url> element to the <urlset> ElementTree object."""
    url_element = ET.SubElement(urlset, "url")
    ET.SubElement(url_element, "loc").text = loc

    if lastmod:
        if isinstance(lastmod, datetime):
            if lastmod.tzinfo is None:
                # Assume UTC if timezone info is missing from DB datetime
                lastmod = lastmod.replace(tzinfo=timezone.utc)
            # Format to YYYY-MM-DD (widely accepted by crawlers)
            formatted_lastmod = lastmod.strftime("%Y-%m-%d")
            ET.SubElement(url_element, "lastmod").text = formatted_lastmod
        elif isinstance(lastmod, str):  # Allow pre-formatted strings
            ET.SubElement(url_element, "lastmod").text = lastmod

    if changefreq:
        ET.SubElement(url_element, "changefreq").text = changefreq
    if priority:
        ET.SubElement(url_element, "priority").text = str(
            priority
        )  # Priority must be a string


def generate_news_sitemap():
    """Generates a news-specific sitemap for better SEO."""
    try:
        urlset = ET.Element(
            "urlset", xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        )

        # Get all visible news articles with SEO data
        news_items = (
            db.session.query(
                News.id, 
                News.updated_at, 
                News.seo_slug,
                News.meta_robots,
                News.is_news,
                News.is_premium,
                News.is_main_news,
                News.read_count,
                News.date,
                News.title,
                News.category_id
            )
            .filter(News.is_visible == True)
            .order_by(News.updated_at.desc())
            .all()
        )
        
        for news_data in news_items:
            news_id, updated_at, seo_slug, meta_robots, is_news, is_premium, is_main_news, read_count, date, title, category_id = news_data
            
            # Skip if explicitly marked as noindex
            if meta_robots and "noindex" in meta_robots.lower():
                continue
                
            try:
                # Use SEO slug if available, otherwise use ID
                if seo_slug:
                    loc = url_for("main.news_detail", news_id=news_id, news_title=seo_slug, _external=True)
                else:
                    loc = url_for("main.news_detail", news_id=news_id, _external=True)
                
                # Determine priority based on content importance
                if is_main_news:
                    priority = 0.9
                elif is_premium:
                    priority = 0.8
                elif read_count and read_count > 100:
                    priority = 0.8
                else:
                    priority = 0.7
                
                # Determine change frequency
                if is_main_news:
                    changefreq = "daily"
                elif is_premium:
                    changefreq = "weekly"
                else:
                    days_old = (datetime.now(timezone.utc) - date).days if date else 365
                    if days_old < 7:
                        changefreq = "daily"
                    elif days_old < 30:
                        changefreq = "weekly"
                    elif days_old < 90:
                        changefreq = "monthly"
                    else:
                        changefreq = "yearly"
                
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"News sitemap: Could not generate URL for news ID {news_id}: {e}"
                )

        sitemap_xml_string = ET.tostring(urlset, encoding="unicode", method="xml")
        final_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + sitemap_xml_string

        return Response(final_xml, mimetype="application/xml")

    except Exception as err:
        current_app.logger.error(
            f"News sitemap generation failed: {err}", exc_info=True
        )
        return Response("Error generating news sitemap", status=500, mimetype="text/plain")


@main_blueprint.route("/sitemap-news.xml")
def sitemap_news():
    """Generates a news-specific sitemap for better SEO."""
    return generate_news_sitemap()


def generate_albums_sitemap():
    """Generates an albums-specific sitemap for better SEO."""
    try:
        urlset = ET.Element(
            "urlset", xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        )

        # Get all visible albums with SEO data
        albums = (
            db.session.query(
                Album.id, 
                Album.updated_at, 
                Album.title,
                Album.is_visible,
                Album.is_archived,
                Album.total_reads,
                Album.created_at,
                Album.is_completed,
                Album.is_hiatus,
                Album.is_premium
            )
            .filter(Album.is_visible == True, Album.is_archived == False)
            .order_by(Album.updated_at.desc())
            .all()
        )
        
        for album_data in albums:
            album_id, updated_at, title, is_visible, is_archived, total_reads, created_at, is_completed, is_hiatus, is_premium = album_data
            
            try:
                # Generate SEO-friendly URL
                safe_album_title = safe_title(title) if title else f"album-{album_id}"
                loc = url_for("main.album_detail", album_id=album_id, album_title=safe_album_title, _external=True)
                
                # Determine priority based on album importance
                if is_premium:
                    priority = 0.9  # Premium albums get highest priority
                elif total_reads and total_reads > 1000:
                    priority = 0.8  # Popular albums get high priority
                elif is_completed:
                    priority = 0.8  # Completed albums get high priority
                else:
                    priority = 0.7  # Regular albums
                
                # Determine change frequency based on album status and popularity
                if is_hiatus:
                    changefreq = "monthly"  # Hiatus albums change less frequently
                elif is_completed:
                    changefreq = "monthly"  # Completed albums change less frequently
                elif is_premium:
                    changefreq = "weekly"  # Premium albums get regular updates
                elif total_reads and total_reads > 500:
                    changefreq = "weekly"  # Popular albums get regular updates
                else:
                    changefreq = "monthly"  # Regular albums
                
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Albums sitemap: Could not generate URL for album ID {album_id}: {e}"
                )

        # Get all visible album chapters
        chapters = (
            db.session.query(
                AlbumChapter.id,
                AlbumChapter.album_id,
                AlbumChapter.chapter_title,
                AlbumChapter.updated_at,
                Album.title.label('album_title'),
                Album.is_visible,
                Album.is_archived,
                Album.is_premium
            )
            .join(Album, AlbumChapter.album_id == Album.id)
            .filter(Album.is_visible == True, Album.is_archived == False)
            .order_by(AlbumChapter.updated_at.desc())
            .all()
        )
        
        for chapter_data in chapters:
            chapter_id, album_id, chapter_title, updated_at, album_title, is_visible, is_archived, is_premium = chapter_data
            
            try:
                # Generate SEO-friendly URLs for both album and chapter titles
                safe_album_title = safe_title(album_title) if album_title else f"album-{album_id}"
                safe_chapter_title = safe_title(chapter_title) if chapter_title else f"chapter-{chapter_id}"
                
                loc = url_for(
                    "main.chapter_reader", 
                    album_id=album_id, 
                    chapter_id=chapter_id, 
                    chapter_title=safe_chapter_title, 
                    _external=True
                )
                
                # Determine priority based on album importance
                if is_premium:
                    priority = 0.8  # Premium album chapters get high priority
                else:
                    priority = 0.7  # Regular chapters
                
                # Chapters typically get updated when new content is added
                changefreq = "weekly" if is_premium else "monthly"
                
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Albums sitemap: Could not generate URL for chapter ID {chapter_id}: {e}"
                )

        sitemap_xml_string = ET.tostring(urlset, encoding="unicode", method="xml")
        final_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + sitemap_xml_string

        return Response(final_xml, mimetype="application/xml")

    except Exception as err:
        current_app.logger.error(
            f"Albums sitemap generation failed: {err}", exc_info=True
        )
        return Response(f"Error generating albums sitemap: {str(err)}", status=500, mimetype="text/plain")


@main_blueprint.route("/sitemap-albums.xml")
def sitemap_albums():
    """Generates an albums-specific sitemap for better SEO."""
    return generate_albums_sitemap()


@main_blueprint.route("/sitemap-index.xml")
def sitemap_index():
    """Generates a sitemap index for large sites with multiple sitemaps."""
    try:
        # Create the root <sitemapindex> element with namespace
        sitemapindex = ET.Element(
            "sitemapindex", xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        )

        # Main sitemap
        sitemap_element = ET.SubElement(sitemapindex, "sitemap")
        ET.SubElement(sitemap_element, "loc").text = url_for("main.sitemap", _external=True)
        ET.SubElement(sitemap_element, "lastmod").text = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        # News-specific sitemap
        sitemap_element = ET.SubElement(sitemapindex, "sitemap")
        ET.SubElement(sitemap_element, "loc").text = url_for("main.sitemap_news", _external=True)
        ET.SubElement(sitemap_element, "lastmod").text = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        # Albums-specific sitemap
        sitemap_element = ET.SubElement(sitemapindex, "sitemap")
        ET.SubElement(sitemap_element, "loc").text = url_for("main.sitemap_albums", _external=True)
        ET.SubElement(sitemap_element, "lastmod").text = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        # Convert to XML string
        sitemap_xml_string = ET.tostring(sitemapindex, encoding="unicode", method="xml")
        final_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + sitemap_xml_string

        return Response(final_xml, mimetype="application/xml")

    except Exception as err:
        current_app.logger.error(
            f"Sitemap index generation failed: {err}", exc_info=True
        )
        return Response("Error generating sitemap index", status=500, mimetype="text/plain")


@main_blueprint.route("/sitemap.xml")
def sitemap():
    """Generates a comprehensive sitemap.xml dynamically with SEO optimization."""
    try:
        # Create the root <urlset> element with namespace
        urlset = ET.Element(
            "urlset", xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        )

        # 1. Static Pages (Publicly accessible HTML pages)
        static_endpoints = [
            ("main.home", "daily", 1.0),  # Homepage - highest priority
            ("main.news", "daily", 0.9),  # Main news listing page
            ("main.videos", "weekly", 0.8),  # Main video listing page
            ("main.gallery", "weekly", 0.7),  # Main gallery listing page
            ("main.utama", "daily", 0.8),  # Utama page
            ("main.about", "monthly", 0.5),  # About page
            ("main.hypes", "weekly", 0.6),  # Hypes page
            ("main.premium", "weekly", 0.6),  # Premium page
        ]

        for endpoint, changefreq, priority in static_endpoints:
            try:
                loc = url_for(endpoint, _external=True)
                add_url_to_sitemap(
                    urlset, loc, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for static endpoint '{endpoint}': {e}"
                )

        # 2. News Detail Pages with SEO optimization
        news_items = (
            db.session.query(
                News.id, 
                News.updated_at, 
                News.seo_slug,
                News.meta_robots,
                News.is_news,
                News.is_premium,
                News.is_main_news,
                News.read_count,
                News.date
            )
            .filter(News.is_visible == True)
            .order_by(News.updated_at.desc())
            .all()
        )
        
        for news_data in news_items:
            news_id, updated_at, seo_slug, meta_robots, is_news, is_premium, is_main_news, read_count, date = news_data
            
            # Skip if explicitly marked as noindex
            if meta_robots and "noindex" in meta_robots.lower():
                continue
                
            try:
                # Use SEO slug if available, otherwise use ID
                if seo_slug:
                    loc = url_for("main.news_detail", news_id=news_id, news_title=seo_slug, _external=True)
                else:
                    loc = url_for("main.news_detail", news_id=news_id, _external=True)
                
                # Determine priority based on content importance
                if is_main_news:
                    priority = 0.9  # Main news gets higher priority
                elif is_premium:
                    priority = 0.8  # Premium content gets higher priority
                elif read_count and read_count > 100:
                    priority = 0.8  # Popular content gets higher priority
                else:
                    priority = 0.7  # Regular content
                
                # Determine change frequency based on content type and age
                if is_main_news:
                    changefreq = "daily"
                elif is_premium:
                    changefreq = "weekly"
                else:
                    # Determine based on how recent the content is
                    days_old = (datetime.now(timezone.utc) - date).days if date else 365
                    if days_old < 7:
                        changefreq = "daily"
                    elif days_old < 30:
                        changefreq = "weekly"
                    elif days_old < 90:
                        changefreq = "monthly"
                    else:
                        changefreq = "yearly"
                
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for news ID {news_id}: {e}"
                )

        # 3. Album Detail Pages with SEO optimization
        albums = (
            db.session.query(
                Album.id, 
                Album.updated_at, 
                Album.title,
                Album.is_visible,
                Album.is_archived,
                Album.total_reads,
                Album.created_at,
                Album.is_completed,
                Album.is_hiatus,
                Album.is_premium
            )
            .filter(Album.is_visible == True, Album.is_archived == False)
            .order_by(Album.updated_at.desc())
            .all()
        )
        
        for album_data in albums:
            album_id, updated_at, title, is_visible, is_archived, total_reads, created_at, is_completed, is_hiatus, is_premium = album_data
            
            try:
                # Generate SEO-friendly URL
                safe_album_title = safe_title(title) if title else f"album-{album_id}"
                loc = url_for("main.album_detail", album_id=album_id, album_title=safe_album_title, _external=True)
                
                # Determine priority based on album importance
                if is_premium:
                    priority = 0.9  # Premium albums get highest priority
                elif total_reads and total_reads > 1000:
                    priority = 0.8  # Popular albums get high priority
                elif is_completed:
                    priority = 0.8  # Completed albums get high priority
                else:
                    priority = 0.7  # Regular albums
                
                # Determine change frequency based on album status and popularity
                if is_hiatus:
                    changefreq = "monthly"  # Hiatus albums change less frequently
                elif is_completed:
                    changefreq = "monthly"  # Completed albums change less frequently
                elif is_premium:
                    changefreq = "weekly"  # Premium albums get regular updates
                elif total_reads and total_reads > 500:
                    changefreq = "weekly"  # Popular albums get regular updates
                else:
                    changefreq = "monthly"  # Regular albums
                
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for album ID {album_id}: {e}"
                )

        # 4. Album Chapter Pages with SEO optimization
        chapters = (
            db.session.query(
                AlbumChapter.id,
                AlbumChapter.album_id,
                AlbumChapter.title,
                AlbumChapter.updated_at,
                Album.title.label('album_title'),
                Album.is_visible,
                Album.is_archived,
                Album.is_premium
            )
            .join(Album, AlbumChapter.album_id == Album.id)
            .filter(Album.is_visible == True, Album.is_archived == False)
            .order_by(AlbumChapter.updated_at.desc())
            .all()
        )
        
        for chapter_data in chapters:
            chapter_id, album_id, chapter_title, updated_at, album_title, is_visible, is_archived, is_premium = chapter_data
            
            try:
                # Generate SEO-friendly URLs for both album and chapter titles
                safe_album_title = safe_title(album_title) if album_title else f"album-{album_id}"
                safe_chapter_title = safe_title(chapter_title) if chapter_title else f"chapter-{chapter_id}"
                
                loc = url_for(
                    "main.chapter_reader", 
                    album_id=album_id, 
                    chapter_id=chapter_id, 
                    chapter_title=safe_chapter_title, 
                    _external=True
                )
                
                # Determine priority based on album importance
                if is_premium:
                    priority = 0.8  # Premium album chapters get high priority
                else:
                    priority = 0.7  # Regular chapters
                
                # Chapters typically get updated when new content is added
                changefreq = "weekly" if is_premium else "monthly"
                
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for chapter ID {chapter_id}: {e}"
                )
        
        for news_data in news_items:
            news_id, updated_at, seo_slug, meta_robots, is_news, is_premium, is_main_news, read_count, date = news_data
            
            # Skip if explicitly marked as noindex
            if meta_robots and "noindex" in meta_robots.lower():
                continue
                
            try:
                # Use SEO slug if available, otherwise use ID
                if seo_slug:
                    loc = url_for("main.news_detail", news_id=news_id, news_title=seo_slug, _external=True)
                else:
                    loc = url_for("main.news_detail", news_id=news_id, _external=True)
                
                # Determine priority based on content importance
                if is_main_news:
                    priority = 0.9  # Main news gets higher priority
                elif is_premium:
                    priority = 0.8  # Premium content gets higher priority
                elif read_count and read_count > 100:
                    priority = 0.8  # Popular content gets higher priority
                else:
                    priority = 0.7  # Regular content
                
                # Determine change frequency based on content type and age
                if is_main_news:
                    changefreq = "daily"
                elif is_premium:
                    changefreq = "weekly"
                else:
                    # Determine based on how recent the content is
                    days_old = (datetime.now(timezone.utc) - date).days if date else 365
                    if days_old < 7:
                        changefreq = "daily"
                    elif days_old < 30:
                        changefreq = "weekly"
                    elif days_old < 90:
                        changefreq = "monthly"
                    else:
                        changefreq = "yearly"
                
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq=changefreq, priority=priority
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for news ID {news_id}: {e}"
                )

        # 5. YouTube Videos
        videos = (
            db.session.query(YouTubeVideo.id, YouTubeVideo.updated_at, YouTubeVideo.youtube_id)
            .filter(YouTubeVideo.is_visible == True)
            .order_by(YouTubeVideo.updated_at.desc())
            .all()
        )
        
        for video_id, updated_at, youtube_id in videos:
            try:
                # Note: This assumes you have a video detail route
                # If not, you might want to link directly to YouTube or create a detail page
                loc = url_for("main.videos", _external=True) + f"#video-{youtube_id}"
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq="weekly", priority=0.6
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for video ID {video_id}: {e}"
                )

        # 6. Images/Gallery Items
        images = (
            db.session.query(Image.id, Image.updated_at, Image.filename)
            .filter(Image.is_visible == True)
            .order_by(Image.updated_at.desc())
            .all()
        )
        
        for image_id, updated_at, filename in images:
            try:
                # Link to gallery page with image focus
                loc = url_for("main.gallery", _external=True) + f"#image-{image_id}"
                add_url_to_sitemap(
                    urlset, loc, lastmod=updated_at, changefreq="monthly", priority=0.5
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for image ID {image_id}: {e}"
                )

        # 7. Category Pages with SEO optimization
        categories = Category.query.order_by(Category.name).all()
        for cat in categories:
            try:
                loc = url_for("main.news", category=cat.name, _external=True)
                # Categories get high priority as they're important for SEO
                add_url_to_sitemap(
                    urlset, loc, changefreq="daily", priority=0.8
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for category '{cat.name}': {e}"
                )

        # 8. Tag Pages with SEO optimization
        tag_strings_tuples = (
            db.session.query(News.tagar)
            .filter(
                and_(News.is_visible == True, News.tagar.isnot(None), News.tagar != "")
            )
            .distinct()
            .all()
        )
        all_tags = set()
        for tag_string_tuple in tag_strings_tuples:
            tag_string = tag_string_tuple[0] if tag_string_tuple[0] else ""
            tags_in_string = [
                tag.strip() for tag in tag_string.split(",") if tag.strip()
            ]
            all_tags.update(tags_in_string)

        for tag_name in sorted(list(all_tags)):
            try:
                loc = url_for("main.news", tag=tag_name, _external=True)
                # Tags get good priority but slightly lower than categories
                add_url_to_sitemap(
                    urlset, loc, changefreq="daily", priority=0.7
                )
            except Exception as e:
                current_app.logger.warning(
                    f"Sitemap: Could not generate URL for tag '{tag_name}': {e}"
                )

        # 9. Pagination pages for main content types (optional, for large sites)
        # News pagination
        total_news = News.query.filter_by(is_visible=True).count()
        if total_news > 20:  # Only add pagination if there are many articles
            total_pages = (total_news + 19) // 20  # Assuming 20 per page
            for page in range(2, min(total_pages + 1, 11)):  # Limit to first 10 pages
                try:
                    loc = url_for("main.news", page=page, _external=True)
                    add_url_to_sitemap(
                        urlset, loc, changefreq="daily", priority=0.6
                    )
                except Exception as e:
                    current_app.logger.warning(
                        f"Sitemap: Could not generate URL for news page {page}: {e}"
                    )

        # Album pagination
        total_albums = Album.query.filter_by(is_visible=True, is_archived=False).count()
        if total_albums > 12:  # Only add pagination if there are many albums
            total_pages = (total_albums + 11) // 12  # Assuming 12 per page
            for page in range(2, min(total_pages + 1, 11)):  # Limit to first 10 pages
                try:
                    loc = url_for("main.albums_list", page=page, _external=True)
                    add_url_to_sitemap(
                        urlset, loc, changefreq="weekly", priority=0.6
                    )
                except Exception as e:
                    current_app.logger.warning(
                        f"Sitemap: Could not generate URL for albums page {page}: {e}"
                    )

        # Video pagination
        total_videos = YouTubeVideo.query.filter_by(is_visible=True).count()
        if total_videos > 6:  # Only add pagination if there are many videos
            total_pages = (total_videos + 5) // 6  # Assuming 6 per page
            for page in range(2, min(total_pages + 1, 6)):  # Limit to first 5 pages
                try:
                    loc = url_for("main.videos", page=page, _external=True)
                    add_url_to_sitemap(
                        urlset, loc, changefreq="weekly", priority=0.5
                    )
                except Exception as e:
                    current_app.logger.warning(
                        f"Sitemap: Could not generate URL for videos page {page}: {e}"
                    )

        # Gallery pagination
        total_images = Image.query.filter_by(is_visible=True).count()
        if total_images > 15:  # Only add pagination if there are many images
            total_pages = (total_images + 14) // 15  # Assuming 15 per page
            for page in range(2, min(total_pages + 1, 6)):  # Limit to first 5 pages
                try:
                    loc = url_for("main.gallery", page=page, _external=True)
                    add_url_to_sitemap(
                        urlset, loc, changefreq="monthly", priority=0.4
                    )
                except Exception as e:
                    current_app.logger.warning(
                        f"Sitemap: Could not generate URL for gallery page {page}: {e}"
                    )

        # Convert the ElementTree to an XML string
        sitemap_xml_string = ET.tostring(urlset, encoding="unicode", method="xml")
        # Add the XML declaration manually as ET doesn't always include it nicely
        final_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + sitemap_xml_string

        return Response(final_xml, mimetype="application/xml")

    except SQLAlchemyError as db_err:
        current_app.logger.error(
            f"Sitemap generation failed due to database error: {db_err}", exc_info=True
        )
        return Response(
            "Error generating sitemap: Database unavailable",
            status=500,
            mimetype="text/plain",
        )
    except Exception as err:
        current_app.logger.error(
            f"Sitemap generation failed due to unexpected error: {err}", exc_info=True
        )
        return Response("Error generating sitemap", status=500, mimetype="text/plain")


@main_blueprint.app_errorhandler(404)
def page_not_found(error):
    return render_template(
        "error.html", code=404, message="Halaman Tidak Ditemukan"
    ), 404


@main_blueprint.app_errorhandler(403)
def forbidden(error):
    return render_template("error.html", code=403, message="Akses Ditolak"), 403


@main_blueprint.app_errorhandler(405)
def method_not_allowed(error):
    return render_template(
        "error.html", code=405, message="Metode Tidak Diizinkan"
    ), 405


@main_blueprint.app_errorhandler(500)
def internal_server_error(error):
    return render_template(
        "error.html", code=500, message="Kesalahan Server Internal"
    ), 500


def save_brand_asset_file(file, field):
    """
    Saves a brand asset file to static/pic/ with a fixed name, replacing the existing file.
    Auto-converts any image to PNG and optimizes size based on the asset type.
    """
    from PIL import Image as PILImage, UnidentifiedImageError
    allowed_names = {
        "logo_header": "logo.png",
        "logo_footer": "logo_footer.png",
        "favicon": "favicon.png",
        "placeholder_image": "placeholder.png",
    }
    
    # Define optimal sizes for each asset type
    optimal_sizes = {
        "logo_header": (300, 80),      # Header logo: 300x80px max
        "logo_footer": (200, 60),      # Footer logo: 200x60px max  
        "favicon": (32, 32),           # Favicon: 32x32px (standard)
        "placeholder_image": (800, 600), # Placeholder: 800x600px max
    }
    
    if field not in allowed_names:
        raise ValueError("Invalid brand asset field")
    
    filename = allowed_names[field]
    file_path = os.path.join("static/pic", filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    
    try:
        # Open and process the image
        img = PILImage.open(file)
        
        # Convert to RGBA for transparency support (important for logos and favicon)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGBA")
        else:
            img = img.convert("RGB")
        
        # Get optimal size for this asset type
        max_width, max_height = optimal_sizes[field]
        
        # Calculate new size while preserving aspect ratio
        img_width, img_height = img.size
        aspect_ratio = img_width / img_height
        
        if aspect_ratio > (max_width / max_height):
            # Image is wider than target ratio
            new_width = max_width
            new_height = int(max_width / aspect_ratio)
        else:
            # Image is taller than target ratio
            new_height = max_height
            new_width = int(max_height * aspect_ratio)
        
        # Resize image if it's larger than optimal size
        if img_width > max_width or img_height > max_height:
            img = img.resize((new_width, new_height), PILImage.Resampling.LANCZOS)
        
        # Special handling for favicon - ensure it's exactly 32x32
        if field == "favicon":
            # Create a 32x32 canvas with transparent background
            favicon = PILImage.new("RGBA", (32, 32), (0, 0, 0, 0))
            
            # Calculate position to center the image
            paste_x = (32 - new_width) // 2
            paste_y = (32 - new_height) // 2
            
            # Paste the resized image onto the canvas
            favicon.paste(img, (paste_x, paste_y), img if img.mode == "RGBA" else None)
            img = favicon
        
        # Save with optimization
        if field == "favicon":
            # Save favicon as ICO format for better browser compatibility
            ico_path = os.path.join("static/pic", "favicon.ico")
            img.save(ico_path, format="ICO", sizes=[(32, 32)])
            # Also save as PNG for fallback
            img.save(file_path, format="PNG", optimize=True)
        else:
            # Save as PNG with optimization
            img.save(file_path, format="PNG", optimize=True)
            
    except UnidentifiedImageError:
        raise ValueError("File is not a valid image")
    except Exception as e:
        raise ValueError(f"Error processing image: {str(e)}")
    
    return filename, file_path
